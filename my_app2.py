from docx import Document 
from docx.shared import Inches
import pyttsx3 

def speak(text):
    pyttsx3.speak(text)

document = Document()

# Profile Piture

document.add_picture(
    'pedrito.jpg', 
    width=Inches(2.0)
    )

# Name, phone number and email details

name = input('What is your name?')
speak('Hello ' + name + ',' ' how are you today?')
phone_number = input('Enter your phone number')
email = input('Enter your email address'+'(Optional)') 
document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)


# About me

document.add_heading('About me')
about_me = input('Tell us about yourself! ')
document.add_paragraph(about_me)

# about_me can be stored either in a variable or via document.add_paragraph(input)('Txt')

# Work experience 

document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter previous or current company')
start_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(start_date + '-' + to_date + '\n').italic = True # '\n' simply indicates that the next paragraph will start in a new line not the same one.

experience_details = input(
    'Describe your experience at ' + company)
p.add_run(experience_details)

# More Experiences
valid = True

while valid:
    
    has_more_expenriences = input(
        'Do you have any more work experience? Yes or No ')
    if has_more_expenriences.lower() == 'yes': 
        p = document.add_paragraph()

        company = input('Enter company ')
        start_date = input('From Date  ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(start_date + '-' + to_date + '\n').italic = True # '\n' simply indicates that the next paragraph will start in a new line not the same one.

        experience_details = input(
            'Describe your experience at ' + company + ' ')
        p.add_run(experience_details)

        #Skills
    else:
        
        
        
        document.add_heading('Skills')
        
        p = document.add_paragraph()

        Skills = input('Tell us one of your skills')
        
        p.add_run(Skills + ' ').bold = True
         
        p.style = 'List Bullet'

        while True:
            has_skill = input(
                'Do you have any more skills? Yes or No')
            if has_skill.lower() == 'yes':
                p = document.add_paragraph()

                Skills = input('Tell us one of your skills')

                p.add_run(Skills + ' ').bold = True
                
                p.style = 'List Bullet'
            elif has_skill == 'no':
                valid = False
                break

#Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'Cv genated by Frank Perez, for business inquiries please email me at frank.nakazona@hotmail.com'
    
        
document.save('cv.docx')



